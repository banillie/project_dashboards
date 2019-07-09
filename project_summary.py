'''Programme that compiles project dashboards/summary sheets.

Input:
1) four quarters worth of data

Output:
1) MS word document in structure of summary sheet / dashboard - with some areas missing, see below.

Supplementary programmes that need to be run to build charts for summary pages. Charts should be built and cut and paste
into dashboards/summary sheets:
1) project_finacial_profile.py . For financial charts
2) milestone_comparison_3_quarters_ind.py . For milestones tables

'''


from docx import Document
from bcompiler.utils import project_data_from_master
import datetime
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.shared import Cm, RGBColor
import difflib


def converting_RAGs(rag):
    if rag == 'Green':
        return 'G'
    elif rag == 'Amber/Green':
        return 'A/G'
    elif rag == 'Amber':
        return 'A'
    elif rag == 'Amber/Red':
        return 'A/R'
    elif rag == 'Red':
        return 'R'
    else:
        return 'None'

def cell_colouring(cell, colour):
    try:
        if colour == 'R':
            colour = parse_xml(r'<w:shd {} w:fill="cb1f00"/>'.format(nsdecls('w')))
        elif colour == 'A/R':
            colour = parse_xml(r'<w:shd {} w:fill="f97b31"/>'.format(nsdecls('w')))
        elif colour == 'A':
            colour = parse_xml(r'<w:shd {} w:fill="fce553"/>'.format(nsdecls('w')))
        elif colour == 'A/G':
            colour = parse_xml(r'<w:shd {} w:fill="a5b700"/>'.format(nsdecls('w')))
        elif colour == 'G':
            colour = parse_xml(r'<w:shd {} w:fill="17960c"/>'.format(nsdecls('w')))

        cell._tc.get_or_add_tcPr().append(colour)

    except TypeError:
        pass

'''function places text into doc highlighing all changes'''
def compare_text_showall(text_1, text_2, doc):
    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.strike = True
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i])

    return doc

'''function places text into doc highlighing new and old text'''
def compare_text_newandold(text_1, text_2, doc):
    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            pass
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i][1:])

    return doc

'''function that compiles the summary sheet'''
def printing(name, dictionary_1, dictionary_2, dictionary_3, dictionary_4, milestone_dict):

    dict_list = [dictionary_1, dictionary_2, dictionary_3, dictionary_4]

    doc = Document()
    print(name)
    heading = str(name)
    name = str(name)
    # TODO: change heading font size
    # todo be able to change text sixe and font
    intro = doc.add_heading(str(heading), 0)
    intro.alignment = 1
    intro.bold = True

    y = doc.add_paragraph()
    a = dictionary_1[name]['Senior Responsible Owner (SRO)']
    if a == None:
        a = 'TBC'
    else:
        a = a
        b = dictionary_1[name]['SRO Phone No.']
        if b == None:
            b = 'TBC'
        else:
            b = b

    y.add_run('SRO name:  ' + str(a) + ',   Tele:  ' + str(b))

    y = doc.add_paragraph()
    a = dictionary_1[name]['Project Director (PD)']
    if a == None:
        a = 'TBC'
    else:
        a = a
        b = dictionary_1[name]['PD Phone No.']
        if b == None:
            b = 'TBC'
        else:
            b = b

    y.add_run('PD name:  ' + str(a) + ',   Tele:  ' + str(b))

    '''Start of table with DCA confidence ratings'''
    table1 = doc.add_table(rows=1, cols=5)
    table1.cell(0, 0).width = Cm(7)

    '''quarter information in top row of table is here'''
    for i, quarter in enumerate(quarter_list):
        table1.cell(0, i+1).text = quarter

    # '''setting row height - partially working'''
    # # todo understand row height better
    # row = table1.rows[0]
    # tr = row._tr
    # trPr = tr.get_or_add_trPr()
    # trHeight = OxmlElement('w:trHeight')
    # trHeight.set(qn('w:val'), str(200))
    # trHeight.set(qn('w:hRule'), 'atLeast')
    # trPr.append(trHeight)

    SRO_conf_table_list = ['SRO DCA', 'Finance DCA', 'Benefits DCA', 'Resourcing DCA', 'Schedule DCA']
    SRO_conf_key_list = ['Departmental DCA', 'SRO Finance confidence', 'SRO Benefits RAG', 'Overall Resource DCA - Now',
                         'SRO Schedule Confidence']

    '''All SRO RAG rating placed in table'''
    for i in range(0, len(dict_list)+1):
        table = doc.add_table(rows=1, cols=5)
        table.cell(0, 0).width = Cm(7)
        table.cell(0, 0).text = SRO_conf_table_list[i]
        for x, dictionary in enumerate(dict_list):
            try:
                rating = converting_RAGs(dictionary[name][SRO_conf_key_list[i]])
                table.cell(0, x + 1).text = rating
                cell_colouring(table.cell(0, x + 1), rating)
            except KeyError:
                table.cell(0, x + 1).text = 'N/A'

    '''DCA Narrative text'''
    doc.add_paragraph()
    y = doc.add_paragraph()
    heading = 'SRO Overall DCA Narrative'
    y.add_run(str(heading)).bold = True

    dca_a = dictionary_1[name]['Departmental DCA Narrative']
    try:
        dca_b = dictionary_2[name]['Departmental DCA Narrative']
    except KeyError:
        dca_b = dca_a

    '''comparing text options'''
    # compare_text_showall(dca_a, dca_b, doc)
    compare_text_newandold(dca_a, dca_b, doc)

    '''Finance section'''
    y = doc.add_paragraph()
    heading = 'Financial information'
    y.add_run(str(heading)).bold = True

    '''Financial Meta data'''
    table1 = doc.add_table(rows=2, cols=5)
    table1.cell(0, 0).text = 'Forecast Whole Life Cost(£m):'
    table1.cell(0, 1).text = 'Percentage Spent:'
    table1.cell(0, 2).text = 'Source of Funding:'
    table1.cell(0, 3).text = 'Nominal or Real figures:'
    table1.cell(0, 4).text = 'Full profile reported:'

    table1.cell(1, 0).text = str(dictionary_1[name]['Total Forecast'])
    a = dictionary_1[name]['Total Forecast']
    b = dictionary_1[name]['Pre 19-20 RDEL Forecast Total']
    if b == None:
        b = 0
    c = dictionary_1[name]['Pre 19-20 CDEL Forecast Total']
    if c == None:
        c = 0
    d = dictionary_1[name]['Pre 19-20 Forecast Non-Gov']
    if d == None:
        d = 0
    e = b + c + d
    try:
        c = round(e / a * 100, 1)
    except ZeroDivisionError:
        c = 0
    table1.cell(1, 1).text = str(c) + '%'
    a = str(dictionary_1[name]['Source of Finance'])
    b = dictionary_1[name]['Other Finance type Description']
    if b == None:
        table1.cell(1, 2).text = a
    else:
        table1.cell(1, 2).text = a + ' ' + str(b)
    table1.cell(1, 3).text = str(dictionary_1[name]['Real or Nominal - Actual/Forecast'])
    table1.cell(1, 4).text = ''

    '''Finance DCA Narrative text'''
    doc.add_paragraph()
    y = doc.add_paragraph()
    heading = 'SRO Finance Narrative'
    y.add_run(str(heading)).bold = True

    narrative = combine_narrtives(name, dictionary_1, gmpp_narrative_keys)
    if narrative == 'NoneNoneNone':
        fin_text = combine_narrtives(name, dictionary_1, bicc_narrative_keys)
    else:
        fin_text = narrative

    compare_text_newandold(fin_text, fin_text, doc)
    #compare_text_showall()

    '''financial chart heading'''  # new
    y = doc.add_paragraph()
    heading = 'Financial Analysis - Cost Profile'
    y.add_run(str(heading)).bold = True
    y = doc.add_paragraph()
    y.add_run('{insert chart}')

    '''milestone section'''
    y = doc.add_paragraph()
    heading = 'Planning information'
    y.add_run(str(heading)).bold = True

    '''Milestone Meta data'''
    table1 = doc.add_table(rows=2, cols=4)
    table1.cell(0, 0).text = 'Project Start Date:'
    table1.cell(0, 1).text = 'Latest Approved Business Case:'
    table1.cell(0, 2).text = 'Start of Operations:'
    table1.cell(0, 3).text = 'Project End Date:'

    key_dates = milestone_dict[name]

    c = key_dates['Start of Project']
    try:
        c = datetime.datetime.strptime(c.isoformat(), '%Y-%M-%d').strftime('%d/%M/%Y')
    except AttributeError:
        c = 'None'

    table1.cell(1, 0).text = str(c)
    table1.cell(1, 1).text = str(dictionary_1[name]['BICC approval point'])

    a = key_dates['Start of Operation']
    try:
        a = datetime.datetime.strptime(a.isoformat(), '%Y-%M-%d').strftime('%d/%M/%Y')
        table1.cell(1, 2).text = str(a)
    except AttributeError:
        table1.cell(1, 2).text = 'None'

    b = key_dates['Project End Date']
    try:
        b = datetime.datetime.strptime(b.isoformat(), '%Y-%M-%d').strftime('%d/%M/%Y')
    except AttributeError:
        b = 'Not reported'
    table1.cell(1, 3).text = str(b)

    # TODO: workout generally styling options for doc, paragraphs and tables

    '''milestone narrative text'''
    doc.add_paragraph()
    y = doc.add_paragraph()
    heading = 'SRO Milestone Narrative'
    y.add_run(str(heading)).bold = True

    mile_dca_a = dictionary_1[name]['Milestone Commentary']
    if mile_dca_a == None:
        mile_dca_a = 'None'

    try:
        mile_dca_b = dictionary_2[name]['Milestone Commentary']
        if mile_dca_b == None:
            mile_dca_b = 'None'
    except KeyError:
        mile_dca_b = mile_dca_a

    # compare_text_showall()
    compare_text_newandold(mile_dca_a, mile_dca_b, doc)

    '''milestone chart heading'''
    y = doc.add_paragraph()
    heading = 'Project reported high-level milestones and schedule changes'
    y.add_run(str(heading)).bold = True
    y = doc.add_paragraph()
    y.add_run('The below table presents all project reported remaining high-level milestones, with six months grace '
              'from close of the current quarter. Milestones are sorted in chronological order. Changes in milestones '
              'dates in comparison to last quarter and one year ago have been calculated and are provided')
    y = doc.add_paragraph()
    y.add_run('{insert chart}')

    return doc

def combine_narrtives(name, dict, key_list):
    output = ''
    for key in key_list:
        output = output + str(dict[name][key])

    return output

def all_milestone_data(master_data):
    upper_dict = {}

    for name in master_data:
        p_data = master_data[name]
        lower_dict = {}
        for i in range(1, 50):
            try:
                try:
                    lower_dict[p_data['Approval MM' + str(i)]] = p_data['Approval MM' + str(i) + ' Forecast / Actual']
                except KeyError:
                    lower_dict[p_data['Approval MM' + str(i)]] = p_data['Approval MM' + str(i) + ' Forecast - Actual']
            except KeyError:
                pass

            try:
                lower_dict[p_data['Assurance MM' + str(i)]] = p_data['Assurance MM' + str(i) + ' Forecast - Actual']
            except:
                pass

        for i in range(18, 67):
            try:
                lower_dict[p_data['Project MM' + str(i)]] = p_data['Project MM' + str(i) + ' Forecast - Actual']
            except:
                pass

        upper_dict[name] = lower_dict

    return upper_dict

quarter_list = ['This Quarter', 'Q4 1819', 'Q3 1819', 'Q2 1819']

gmpp_narrative_keys = ['Project Costs Narrative', 'Cost comparison with last quarters cost - narrative',
                  'Cost comparison within this quarters cost - narrative']

bicc_narrative_keys = ['Project Costs Narrative RDEL', 'Project Costs Narrative CDEL']

current_Q_dict = project_data_from_master('C:\\Users\\Standalone\\Will\\masters folder\\core data\\Hs2_NPR_Q1_1918_draft.xlsx')
last_Q_dict = project_data_from_master('C:\\Users\\Standalone\\Will\\masters folder\\core data\\master_4_2018.xlsx')
Q2_ago_dict = project_data_from_master('C:\\Users\\Standalone\\Will\\masters folder\\core data\\master_3_2018.xlsx')
Q3_ago_dict = project_data_from_master('C:\\Users\\Standalone\\Will\\masters folder\\core data\\master_2_2018.xlsx')

current_Q_list = list(current_Q_dict.keys())
#current_Q_list = ['High Speed Rail Programme (HS2)']
# current_Q_list.remove('Commercial Vehicle Services (CVS)')

milestones = all_milestone_data(current_Q_dict)

for x in current_Q_list:
    a = printing(x, current_Q_dict, last_Q_dict, Q2_ago_dict, Q3_ago_dict, milestones)
    a.save('C://Users//Standalone//Will//Q1_1920_{}_overview_test.docx'.format(x))

